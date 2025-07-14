from io import BytesIO
from docx import Document as DocxDocument

def parse_stream(filename: str, content: bytes) -> str:

    if filename.lower().endswith(".txt"):
        return content.decode("utf-8", errors="ignore").strip()

    elif filename.lower().endswith(".docx"):
        try:
            doc = DocxDocument(BytesIO(content))
            texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            print(f"docx paragraphs: {len(doc.paragraphs)}, non-empty: {len(texts)}")
            return "\n".join(texts)
        except Exception as e:
            raise ValueError(f"Failed to parse .docx: {e}")

    else:
        raise ValueError(f"Unsupported file type: {filename}")
